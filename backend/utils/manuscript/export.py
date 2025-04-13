import os
import io
import base64
import re
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
import markdown
import weasyprint

from app.models.schemas import Manuscript, TableData, FigureData

def format_paragraphs(content: str) -> str:
    """
    Format text content into HTML paragraphs.
    
    Args:
        content: Text content to format
        
    Returns:
        str: HTML formatted paragraphs
    """
    if not content:
        return "<p>No content available.</p>"
    
    paragraphs = re.split(r'\n\s*\n', content)
    
    formatted = ""
    for paragraph in paragraphs:
        clean_paragraph = re.sub(r'\n', ' ', paragraph.strip())
        if clean_paragraph:
            formatted += f"<p>{clean_paragraph}</p>\n"
    
    return formatted

def format_references(references: List[str]) -> str:
    """
    Format references as HTML.
    
    Args:
        references: List of reference strings
        
    Returns:
        str: HTML formatted references
    """
    if not references:
        return "<p>No references available.</p>"
    
    formatted = ""
    for i, ref in enumerate(references):
        formatted += f'<div class="reference">{i+1}. {ref}</div>\n'
    
    return formatted

def export_tables_to_docx(doc, tables):
    """
    Add tables to a Word document.
    
    Args:
        doc: Document object to add tables to
        tables: List of TableData objects
        
    Returns:
        Document: Updated document with tables
    """
    if not tables:
        return doc
        
    doc.add_heading("Tables", level=2)
    
    for i, table_data in enumerate(tables):
        doc.add_heading(f"Table {i+1}: {table_data.title}", level=3)
        table = doc.add_table(rows=len(table_data.rows)+1, cols=len(table_data.headers))
        table.style = 'Table Grid'
        
        for i, header in enumerate(table_data.headers):
            table.cell(0, i).text = header
            
        for i, row in enumerate(table_data.rows):
            for j, cell in enumerate(row):
                table.cell(i+1, j).text = cell
                
        caption = doc.add_paragraph(table_data.caption)
        caption.italic = True
        
        doc.add_paragraph("")  # Add space after table
        
    return doc

def export_figures_to_docx(doc, figures):
    """
    Add figures to a Word document.
    
    Args:
        doc: Document object to add figures to
        figures: List of FigureData objects
        
    Returns:
        Document: Updated document with figures
    """
    if not figures:
        return doc
        
    doc.add_heading("Figures", level=2)
    
    for i, figure in enumerate(figures):
        doc.add_heading(f"Figure {i+1}: {figure.title}", level=3)
        
        placeholder = doc.add_paragraph(f"[Figure placeholder: {figure.type} chart - {figure.subtype or 'general'}]")
        
        caption = doc.add_paragraph(figure.caption)
        caption.italic = True
        
        doc.add_paragraph("")  # Add space after figure
        
    return doc

def export_to_docx(manuscript: Manuscript, output_path: str) -> str:
    """
    Export the manuscript to a Word document.
    
    Args:
        manuscript: Manuscript to export
        output_path: Path to save the document
        
    Returns:
        str: Path to the saved document
    """
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    title = doc.add_heading(manuscript.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading(manuscript.abstract.title, level=2)
    doc.add_paragraph(manuscript.abstract.content)
    
    doc.add_heading(manuscript.introduction.title, level=2)
    doc.add_paragraph(manuscript.introduction.content)
    
    doc.add_heading(manuscript.methods.title, level=2)
    doc.add_paragraph(manuscript.methods.content)
    
    doc.add_heading(manuscript.results.title, level=2)
    doc.add_paragraph(manuscript.results.content)
    
    doc.add_heading(manuscript.discussion.title, level=2)
    doc.add_paragraph(manuscript.discussion.content)
    
    doc.add_heading("References", level=2)
    for ref in manuscript.references:
        doc.add_paragraph(ref)
    
    if manuscript.tables:
        doc = export_tables_to_docx(doc, manuscript.tables)
        
    if manuscript.figures:
        doc = export_figures_to_docx(doc, manuscript.figures)
    
    doc.add_paragraph(f"Word Count: {manuscript.word_count}")
    
    doc.save(output_path)
    
    return output_path

def export_to_pdf(manuscript: Manuscript, output_path: str) -> str:
    """
    Export the manuscript to a PDF document using WeasyPrint with improved formatting.
    
    Args:
        manuscript: Manuscript to export
        output_path: Path to save the document
        
    Returns:
        str: Path to the saved document
    """
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>{manuscript.title}</title>
        <style>
            @page {{
                size: letter;
                margin: 1in;
            }}
            body {{
                font-family: 'Times New Roman', Times, serif;
                font-size: 12pt;
                line-height: 1.5;
                counter-reset: page;
            }}
            h1 {{
                text-align: center;
                font-size: 16pt;
                font-weight: bold;
                margin-bottom: 24pt;
            }}
            h2 {{
                font-size: 14pt;
                font-weight: bold;
                margin-top: 24pt;
                margin-bottom: 12pt;
                page-break-after: avoid;
            }}
            h3 {{
                font-size: 12pt;
                font-weight: bold;
                margin-top: 18pt;
                margin-bottom: 6pt;
            }}
            p {{
                margin-bottom: 12pt;
                text-align: justify;
            }}
            .abstract {{
                margin-bottom: 24pt;
            }}
            .references {{
                margin-top: 24pt;
            }}
            .reference {{
                text-indent: -24pt;
                padding-left: 24pt;
                margin-bottom: 6pt;
            }}
            .word-count {{
                margin-top: 24pt;
                text-align: right;
                font-style: italic;
            }}
            .page-break {{
                page-break-before: always;
            }}
            table {{
                width: 100%;
                border-collapse: collapse;
                margin-bottom: 12pt;
            }}
            table, th, td {{
                border: 1px solid black;
            }}
            th, td {{
                padding: 6pt;
                text-align: left;
            }}
            th {{
                background-color: #f2f2f2;
            }}
            .table-caption, .figure-caption {{
                font-style: italic;
                text-align: center;
                margin-bottom: 18pt;
            }}
            .figure-placeholder {{
                width: 100%;
                height: 300px;
                background-color: #f2f2f2;
                display: flex;
                justify-content: center;
                align-items: center;
                margin-bottom: 12pt;
                border: 1px solid #ccc;
            }}
            .table-container, .figure-container {{
                margin-top: 24pt;
                margin-bottom: 24pt;
            }}
        </style>
    </head>
    <body>
        <h1>{manuscript.title}</h1>
        
        <h2>{manuscript.abstract.title}</h2>
        <div class="abstract">
            {format_paragraphs(manuscript.abstract.content)}
        </div>
        
        <h2>{manuscript.introduction.title}</h2>
        <div>
            {format_paragraphs(manuscript.introduction.content)}
        </div>
        
        <h2>{manuscript.methods.title}</h2>
        <div>
            {format_paragraphs(manuscript.methods.content)}
        </div>
        
        <h2>{manuscript.results.title}</h2>
        <div>
            {format_paragraphs(manuscript.results.content)}
        </div>
        
        <h2>{manuscript.discussion.title}</h2>
        <div>
            {format_paragraphs(manuscript.discussion.content)}
        </div>
        
        <h2>References</h2>
        <div class="references">
            {format_references(manuscript.references)}
        </div>
        
        <!-- Add tables if they exist -->
        {"""
        <h2>Tables</h2>
        """ + "".join([f"""
        <div class="table-container">
            <h3>Table {i+1}: {table.title}</h3>
            <table class="manuscript-table">
                <thead>
                    <tr>
                        {"".join([f"<th>{header}</th>" for header in table.headers])}
                    </tr>
                </thead>
                <tbody>
                    {"".join([f"""
                    <tr>
                        {"".join([f"<td>{cell}</td>" for cell in row])}
                    </tr>
                    """ for row in table.rows])}
                </tbody>
            </table>
            <p class="table-caption">{table.caption}</p>
        </div>
        """ for i, table in enumerate(manuscript.tables)]) if manuscript.tables else ""}
        
        <!-- Add figures if they exist -->
        {"""
        <h2>Figures</h2>
        """ + "".join([f"""
        <div class="figure-container">
            <h3>Figure {i+1}: {figure.title}</h3>
            <div class="figure-placeholder">
                [{figure.type.capitalize()} chart: {figure.subtype or "General"}]
            </div>
            <p class="figure-caption">{figure.caption}</p>
        </div>
        """ for i, figure in enumerate(manuscript.figures)]) if manuscript.figures else ""}
        
        <div class="word-count">
            Word Count: {manuscript.word_count}
        </div>
    </body>
    </html>
    """
    
    try:
        pdf = weasyprint.HTML(string=html).write_pdf()
        
        with open(output_path, 'wb') as f:
            f.write(pdf)
        
        return output_path
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        try:
            options = {
                'page-size': 'Letter',
                'margin-top': '1in',
                'margin-right': '1in',
                'margin-bottom': '1in',
                'margin-left': '1in',
                'encoding': 'UTF-8',
                'quiet': ''
            }
            pdfkit.from_string(html, output_path, options=options)
            return output_path
        except Exception as pdfkit_error:
            print(f"Fallback PDF generation also failed: {str(pdfkit_error)}")
            with open(output_path, 'wb') as f:
                f.write(b"%PDF-1.4\n1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n2 0 obj\n<</Type/Pages/Kids[3 0 R]/Count 1>>\nendobj\n3 0 obj\n<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>/Contents 4 0 R>>\nendobj\n4 0 obj\n<</Length 21>>\nstream\nBT\n/F1 12 Tf\n100 700 Td\n(Error generating PDF. Please try downloading as Word or Markdown instead.) Tj\nET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f\n0000000009 00000 n\n0000000056 00000 n\n0000000111 00000 n\n0000000212 00000 n\ntrailer\n<</Size 5/Root 1 0 R>>\nstartxref\n284\n%%EOF")
            return output_path

def manuscript_to_markdown(manuscript: Manuscript) -> str:
    """
    Convert a manuscript to Markdown format.
    
    Args:
        manuscript: Manuscript to convert
        
    Returns:
        str: Markdown representation of the manuscript
    """
    md = f"# {manuscript.title}\n\n"
    
    md += f"## {manuscript.abstract.title}\n\n{manuscript.abstract.content}\n\n"
    
    md += f"## {manuscript.introduction.title}\n\n{manuscript.introduction.content}\n\n"
    
    md += f"## {manuscript.methods.title}\n\n{manuscript.methods.content}\n\n"
    
    md += f"## {manuscript.results.title}\n\n{manuscript.results.content}\n\n"
    
    md += f"## {manuscript.discussion.title}\n\n{manuscript.discussion.content}\n\n"
    
    md += "## References\n\n"
    for ref in manuscript.references:
        md += f"{ref}\n\n"
    
    if manuscript.tables:
        md += "## Tables\n\n"
        for i, table in enumerate(manuscript.tables):
            md += f"### Table {i+1}: {table.title}\n\n"
            
            md += "| " + " | ".join(table.headers) + " |\n"
            md += "| " + " | ".join(["---"] * len(table.headers)) + " |\n"
            
            for row in table.rows:
                md += "| " + " | ".join(row) + " |\n"
            
            md += f"\n*{table.caption}*\n\n"
    
    if manuscript.figures:
        md += "## Figures\n\n"
        for i, figure in enumerate(manuscript.figures):
            md += f"### Figure {i+1}: {figure.title}\n\n"
            
            md += f"[{figure.type.capitalize()} chart: {figure.subtype or 'General'}]\n\n"
            
            md += f"*{figure.caption}*\n\n"
    
    md += f"Word Count: {manuscript.word_count}\n"
    
    return md

def markdown_to_html(md: str) -> str:
    """
    Convert Markdown to HTML.
    
    Args:
        md: Markdown content
        
    Returns:
        str: HTML content
    """
    html = markdown.markdown(md, extensions=['tables'])
    
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>Manuscript</title>
    </head>
    <body>
        {html}
    </body>
    </html>
    """

def export_to_markdown(manuscript: Manuscript, output_path: str) -> str:
    """
    Export the manuscript to a Markdown file.
    
    Args:
        manuscript: Manuscript to export
        output_path: Path to save the document
        
    Returns:
        str: Path to the saved document
    """
    md = manuscript_to_markdown(manuscript)
    
    with open(output_path, 'w') as f:
        f.write(md)
    
    return output_path
